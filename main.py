import csv
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn


#
# Customer = [
#     "301 Reid Dr, Durham, NC 27705", '301 Reid Dr', 'Durham', 'NC', '27705',
#     '104603', 'Shirl', 'Thompson', '301 REID DR, DURHAM, NC, 27705']


# # Construct the heading with the specified format, avoiding repetition
# heading = (
#     f"{Customer[6]} {Customer[7]},\n"  # First Name Last Name
#     f"{Customer[1]}\n"  # Primary Mailing Address
#     f"{Customer[2]} {Customer[3]}, {Customer[4]}\n\n"  # City, State, ZIP (City in proper case)
#     f"Dear {Customer[6]},\n"  # Salutation
# )


def read_csv_file(file_path):
    data_list = []
    with open(file_path, mode='r', newline='', encoding='utf-8') as file:
        csv_reader = csv.reader(file)
        next(csv_reader)  # Skip the header row
        for row in csv_reader:
            data_list.append(row)
        return data_list


def create_letter(doc, customer):
    # Assuming the customer list contains the fields in the following order:
    # [full_address, address_line_1, city, state, zipcode, parcel_id, first_name, last_name, primary_mailing_address]
    full_address = customer[0]
    address_line_1 = customer[1]
    city = customer[2]
    state = customer[3]
    zipcode = customer[4]
    first_name = customer[6]
    last_name = customer[7]
    primary_mailing_address = customer[8]

    heading = (
        f"{first_name} {last_name},\n"  # First Name Last Name
        f"{primary_mailing_address}\n"  # Primary Mailing Address
        f"{city} {state} {zipcode}\n\n"  # City, State, ZIP
        f"Dear {first_name},\n"

    )

    body = (
        f"I am writing you because I am interested in purchasing your property at {full_address}.\n"
        f"I live nearby in Apex, and I purchase properties throughout the state of North Carolina. "
        f"Currently, I am working on projects in the neighborhood which is how I learned of your property.\n\n"

        f"Please call me at (919) 452-9700 so I can tell you how we can buy your property quickly, with no hassle, "
        f"and most importantly, at a fair price. I purchase the house directly from you so there are no real estate "
        f"commissions or extra fees from your pocket. Furthermore, I always make cash offers, so you never have to "
        f"worry"
        f"about delays resulting from a buyer trying to qualify for a mortgage and wasting your valuable time.\n\n"

        f"I'll be glad to personally tell you more about how we can work together. All calls are completely "
        f"confidential"
        f"and there is no obligation. Even if you're not interested in selling at this time, call (919) 452-9700 so I "
        f"can"
        f"help you determine the value of your property. You should also keep this letter for future reference since "
        f"things"
        f"in life tend to change unexpectedly.\n\n"

        f"Thank you for your time and consideration and I’m looking forward to speaking with you!\n\n"

        f"Sincerely,")
    signature = f"Kayla Morabito"
    ending = (
        f"(919) 452-9700\n"

        f"P.S If you would like to email me information and pictures of your property, please send them to "
        f"kkmorabito@aol.com"
    )

    doc.add_paragraph(heading)
    doc.add_paragraph(body)
    # Adding a paragraph with the desired text
    paragraph = doc.add_paragraph()

    # Adding a run within the paragraph and setting the font
    run = paragraph.add_run(signature)

    # Set the font for the run
    run.font.name = 'Charm'  # Set the font name

    # For compatibility with non-Word software, set it in the XML
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Charm')
    rPr.append(rFonts)

    # You can also set the font size, bold, italic, etc.
    run.font.size = Pt(12)  # Example to set font size to 12pt
    # run.bold = True  # To set the text to bold
    doc.add_paragraph(ending)
    doc.add_page_break()


def export_to_csv(data_list, file_path):
    with open(file_path, mode='w', newline='', encoding='utf-8') as file:
        csv_writer = csv.writer(file)
        csv_writer.writerows(data_list)


def remove_duplicates(data_list):
    unique_data = []
    for i in data_list:
        if i not in unique_data:
            unique_data.append(i)
    return unique_data


def remove_blank_cells(data_list):
    cleaned_data = []
    for i in data_list:
        cleaned_row = [cell for cell in i if cell.strip() != ""]
        cleaned_data.append(cleaned_row)
    return cleaned_data


docx_file_path = 'Output.docx'
csv_file_path = 'MiniTest.csv'
new_csv_file_path = 'CleanedMiniTest.csv'


def process_csv_and_generate_doc(template_path, output_doc_path):
    doc_created = False
    try:
        doc = Document(output_doc_path)  # open an existing document
    except Exception as e:
        print(f"Document not found, creating a new one: {e}")
        doc = Document()  # or create a new document
        doc_created = True

    # Read the CSV file and print the data
    data = read_csv_file(template_path)
    if not data:
        print("No data found in the CSV file.")
    else:
        data = remove_duplicates(data)
        data = remove_blank_cells(data)

        # For each customer in the list, create a letter and add it to the document
        for customer in data:
            create_letter(doc, customer)

        doc.save(output_doc_path)  # save the updated document

        if doc_created:
            print(f"New document created and saved to {output_doc_path}")
        else:
            print(f"Document updated and saved to {output_doc_path}")

    # Export the cleaned data to a new CSV file
    export_to_csv(data, 'CleanedMiniTest.csv')
    print("Exported cleaned data to CleanedMiniTest.csv")


if __name__ == "__main__":
    template_path = 'MiniTest.csv'  # input CSV file
    output_doc_path = 'Output.docx'  # output Word document
    process_csv_and_generate_doc(template_path, output_doc_path)

# If import csv is not in the same format as the customer list, the code will need to be adjusted accordingly.
# The code assumes that the customer list contains the fields in the following order:
# associated_property_address_full, associated_property_address_line_1, associated_property_address_line_2
# associated_property_address_city, associated_property_address_state, associated_property_address_zipcode,
# associated_parcel_id, first_name,	last_name, primary_mailing_address, primary_mailing_city, primary_mailing_state,
# primary_mailing_zip

# When reading a csv write a data field checker to see if the header line cells match the expected fields
# To compare headers we will have static strings to compare against.
